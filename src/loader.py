# === src/loader.py ===
# === src/loader.py ===

import os
from langchain.schema import Document
from langchain_community.document_loaders import (
    PyPDFLoader,
    TextLoader,
    UnstructuredMarkdownLoader,
    UnstructuredHTMLLoader,
    UnstructuredEPubLoader,
    CSVLoader,
)
from docx import Document as DocxDocument
import pandas as pd
from odf import text, teletype
from odf.opendocument import load as load_odt
from src.translator import translate_to_english


SUPPORTED_EXTENSIONS = [
    ".pdf",
    ".txt",
    ".md",
    ".html",
    ".htm",
    ".csv",
    ".docx",
    ".epub",
    ".odt",
]

#translate documents in italian
def make_translated_document(text: str, metadata: dict):
    translated = translate_to_english(text)

    new_metadata = dict(metadata)  # copia sicura
    new_metadata.update({
        "lang": "en",
        "original_lang": "it",
        "original_text": text
    })

    return Document(
        page_content=translated,
        metadata=new_metadata
    )

def load_docx_file(path: str):
    """Load docx file and read header footer
    """
    try:
        doc = DocxDocument(path)
    except Exception as e:
        print(f"❌ Error opening DOCX {path}: {e}")
        return []

    parts = []

    # paragrafi
    for p in doc.paragraphs:
        if p.text and p.text.strip():
            parts.append(p.text.strip())

    # tabelle (celle)
    try:
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text.strip()
                    if text:
                        parts.append(text)
    except Exception as e:
        # no blocking, continue even if there are no tables
        print(f"⚠️ Error reading tables in {path}: {e}")

    # if there are header footer
    try:
        if doc.sections:
            for sec in doc.sections:
                hdr = sec.header
                ftr = sec.footer
                if hdr and hdr.paragraphs:
                    for p in hdr.paragraphs:
                        if p.text and p.text.strip():
                            parts.append(p.text.strip())
                if ftr and ftr.paragraphs:
                    for p in ftr.paragraphs:
                        if p.text and p.text.strip():
                            parts.append(p.text.strip())
    except Exception:
        # ignore problems in the section
        pass

    full_text = "\n\n".join(parts).strip()
    if not full_text:
        print(f"⚠️ DOCX {path} loaded but empty.")
        return []

    metadata = {"source": path}
    return [make_translated_document(full_text, metadata)]



def load_csv_file(path: str):
    """Load csv as chained"""
    df = pd.read_csv(path, encoding="utf-8", engine="python")
    text = df.to_string()
    return [make_translated_document(text, {"source": path})]

def load_odt_file(path: str):
    """Load odt as a single document"""
    doc = load_odt(path)
    all_text = []
    for elem in doc.getElementsByType(text.P):
        all_text.append(teletype.extractText(elem))
    full_text = "\n".join(all_text)
    return [make_translated_document(full_text, {"source": path})]

def load_single_file(path: str):
    """Return a langChain documents for each different type of file."""
    ext = os.path.splitext(path)[1].lower()

    if ext == ".pdf":
        docs = PyPDFLoader(path).load()
        return [
            make_translated_document(d.page_content, d.metadata)
            for d in docs
        ]

    elif ext == ".txt":
        docs = TextLoader(path, encoding="utf-8").load()
        return [
            make_translated_document(d.page_content, d.metadata)
            for d in docs
        ]

    elif ext == ".md":
        docs = UnstructuredMarkdownLoader(path).load()
        return [
            make_translated_document(d.page_content, d.metadata)
            for d in docs
        ]

    elif ext in [".html", ".htm"]:
        docs = UnstructuredHTMLLoader(path).load()
        return [
            make_translated_document(d.page_content, d.metadata)
            for d in docs
        ]

    elif ext == ".epub":
        docs = UnstructuredEPubLoader(path).load()
        return [
            make_translated_document(d.page_content, d.metadata)
            for d in docs
        ]

    elif ext == ".csv":
        return load_csv_file(path)

    elif ext == ".docx":
        return load_docx_file(path)

    elif ext == ".odt":
        return load_odt_file(path)

    else:
        print(f"❌ Formato non supportato: {ext}")
        return []


def load_pdf_files(data_path):
    """
    Load every file inside this function
    """
    documents = []

    for filename in os.listdir(data_path):
        path = os.path.join(data_path, filename)

        if not os.path.isfile(path):
            continue

        ext = os.path.splitext(filename)[1].lower()

        if ext in SUPPORTED_EXTENSIONS:
            print(f"📄 Caricamento file: {filename}")
            docs = load_single_file(path)
            documents.extend(docs)
        else:
            print(f"⚠️ Ignorato (formato non supportato): {filename}")

    print(f"📚 Totale documenti caricati: {len(documents)}")
    return documents





